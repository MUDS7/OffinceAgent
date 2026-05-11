from __future__ import annotations

from base64 import b64decode, b64encode
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Literal
from urllib.parse import quote

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.excel_commands import (
    ExcelCommandsResponse,
    ExcelExecuteRequest,
    ExcelExecuteResponse,
    execute_excel_command,
    get_excel_commands,
)


app = FastAPI(title="OfficeAgent Document Service", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_private_network=True,
)


class HealthResponse(BaseModel):
    status: str
    service: str


class AnalyzeResponse(BaseModel):
    filename: str
    extension: str
    size_bytes: int
    sha256: str
    text_preview: str
    warnings: list[str]


class DocxParagraphBlock(BaseModel):
    id: str
    type: Literal["paragraph"] = "paragraph"
    text: str
    style: str | None = None
    style_id: str | None = None
    alignment: str | None = None


class DocxTableCell(BaseModel):
    id: str
    text: str
    alignment: str | None = None


class DocxTableBlock(BaseModel):
    id: str
    type: Literal["table"] = "table"
    rows: list[list[DocxTableCell]]


class DocxImageBlock(BaseModel):
    id: str
    type: Literal["image"] = "image"
    filename: str = "image"
    content_type: str = "application/octet-stream"
    data_url: str
    alt_text: str | None = None
    width_emu: int | None = None
    height_emu: int | None = None
    alignment: str | None = None


DocxBlock = DocxParagraphBlock | DocxTableBlock | DocxImageBlock


class DocxParseResponse(BaseModel):
    filename: str
    blocks: list[DocxBlock]
    text_preview: str
    warnings: list[str]


class DocxRenderRequest(BaseModel):
    filename: str = "document.docx"
    blocks: list[DocxBlock]


@app.get("/health", response_model=HealthResponse)
def health() -> HealthResponse:
    return HealthResponse(status="ok", service="document-service")


@app.post("/documents/analyze", response_model=AnalyzeResponse)
async def analyze_document(file: UploadFile = File(...)) -> AnalyzeResponse:
    content = await file.read()
    filename = file.filename or "untitled"
    extension = Path(filename).suffix.lower()
    warnings: list[str] = []
    text = extract_text(filename, extension, content, warnings)

    return AnalyzeResponse(
        filename=filename,
        extension=extension.lstrip(".") or "unknown",
        size_bytes=len(content),
        sha256=sha256(content).hexdigest(),
        text_preview=text[:4000] if text else "",
        warnings=warnings,
    )


@app.post("/docx/parse", response_model=DocxParseResponse)
async def parse_docx(file: UploadFile = File(...)) -> DocxParseResponse:
    content = await file.read()
    filename = file.filename or "untitled.docx"
    warnings: list[str] = []
    blocks = extract_docx_blocks(content, warnings)
    text_preview = "\n".join(get_docx_block_text(block) for block in blocks).strip()

    return DocxParseResponse(
        filename=filename,
        blocks=blocks,
        text_preview=text_preview[:4000],
        warnings=warnings,
    )


@app.post("/docx/render")
def render_docx(request: DocxRenderRequest) -> StreamingResponse:
    try:
        content = build_docx_bytes(request.blocks)
    except Exception as exc:  # pragma: no cover - external document writer boundary
        raise HTTPException(
            status_code=500,
            detail=(
                f"DOCX 生成失败（文件：{Path(request.filename or 'document.docx').name}；"
                f"{summarize_docx_blocks(request.blocks)}）：{type(exc).__name__}: {exc}"
            ),
        ) from exc

    filename = Path(request.filename or "document.docx").name
    encoded_filename = quote(filename)
    return StreamingResponse(
        BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"},
    )


@app.get("/excel/commands", response_model=ExcelCommandsResponse)
def list_excel_commands() -> ExcelCommandsResponse:
    return get_excel_commands()


@app.post("/excel/execute", response_model=ExcelExecuteResponse)
def run_excel_command(request: ExcelExecuteRequest) -> ExcelExecuteResponse:
    return execute_excel_command(request)


def extract_text(filename: str, extension: str, content: bytes, warnings: list[str]) -> str:
    if extension in {".txt", ".md", ".csv", ".json"}:
        return decode_text(content, warnings)

    if extension == ".pdf":
        return extract_pdf_text(content, warnings)

    if extension == ".docx":
        return extract_docx_text(content, warnings)

    warnings.append(f"{filename} 的类型暂未配置文本抽取器")
    return ""


def decode_text(content: bytes, warnings: list[str]) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    warnings.append("无法按常见编码解码文本")
    return ""


def extract_pdf_text(content: bytes, warnings: list[str]) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages[:10]]
        if len(reader.pages) > 10:
            warnings.append("仅预览 PDF 前 10 页文本")
        return "\n\n".join(pages).strip()
    except Exception as exc:  # pragma: no cover - external parser boundary
        warnings.append(f"PDF 解析失败: {exc}")
        return ""


def extract_docx_text(content: bytes, warnings: list[str]) -> str:
    try:
        blocks = extract_docx_blocks(content, warnings)
        return "\n".join(get_docx_block_text(block) for block in blocks).strip()
    except Exception as exc:  # pragma: no cover - external parser boundary
        warnings.append(f"DOCX 解析失败: {exc}")
        return ""


def extract_docx_blocks(content: bytes, warnings: list[str]) -> list[DocxBlock]:
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(BytesIO(content))
        blocks: list[DocxBlock] = []
        paragraph_index = 0
        table_index = 0
        body = document._body

        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, body)
                text = paragraph.text
                style = paragraph.style.name if paragraph.style else None
                style_id = paragraph.style.style_id if paragraph.style else None
                alignment = get_paragraph_alignment(paragraph)
                if text.strip():
                    blocks.append(
                        DocxParagraphBlock(
                            id=f"p-{paragraph_index}",
                            text=text,
                            style=style,
                            style_id=style_id,
                            alignment=alignment,
                        )
                    )
                blocks.extend(extract_docx_paragraph_images(paragraph, f"p-{paragraph_index}", alignment))
                paragraph_index += 1
                continue

            if isinstance(child, CT_Tbl):
                table = Table(child, body)
                rows: list[list[DocxTableCell]] = []
                for row_index, row in enumerate(table.rows):
                    cells = [
                        DocxTableCell(
                            id=f"t-{table_index}-r-{row_index}-c-{cell_index}",
                            text=cell.text,
                            alignment=get_cell_alignment(cell),
                        )
                        for cell_index, cell in enumerate(row.cells)
                    ]
                    rows.append(cells)

                if rows:
                    blocks.append(
                        DocxTableBlock(
                            id=f"t-{table_index}",
                            rows=rows,
                        )
                    )
                table_index += 1

        if not blocks:
            warnings.append("DOCX 中未提取到可显示文本")

        return blocks
    except Exception as exc:  # pragma: no cover - external parser boundary
        warnings.append(f"DOCX 解析失败: {exc}")
        return []


def get_paragraph_alignment(paragraph) -> str | None:
    alignment = paragraph.alignment
    if alignment is None and paragraph.style:
        alignment = paragraph.style.paragraph_format.alignment
    if alignment is None:
        return None
    return alignment.name.lower().replace("_", "-")


def get_cell_alignment(cell) -> str | None:
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            return get_paragraph_alignment(paragraph)
    return get_paragraph_alignment(cell.paragraphs[0]) if cell.paragraphs else None


def extract_docx_paragraph_images(paragraph, paragraph_id: str, alignment: str | None) -> list[DocxImageBlock]:
    from docx.oxml.ns import qn

    images: list[DocxImageBlock] = []
    image_index = 0
    for run_index, run in enumerate(paragraph.runs):
        blips = run._element.xpath(".//a:blip")
        extents = run._element.xpath(".//wp:extent")
        names = run._element.xpath(".//pic:cNvPr")

        for blip_index, blip in enumerate(blips):
            relationship_id = blip.get(qn("r:embed"))
            if not relationship_id:
                continue

            part = paragraph.part.related_parts.get(relationship_id)
            if part is None:
                continue

            extent = extents[min(blip_index, len(extents) - 1)] if extents else None
            name_node = names[min(blip_index, len(names) - 1)] if names else None
            filename = Path(str(part.partname)).name or f"{paragraph_id}-image-{image_index + 1}"
            content_type = getattr(part, "content_type", None) or "application/octet-stream"
            alt_text = None
            if name_node is not None:
                alt_text = name_node.get("descr") or name_node.get("name")

            images.append(
                DocxImageBlock(
                    id=f"{paragraph_id}-img-{run_index}-{image_index}",
                    filename=filename,
                    content_type=content_type,
                    data_url=f"data:{content_type};base64,{b64encode(part.blob).decode('ascii')}",
                    alt_text=alt_text,
                    width_emu=int(extent.cx) if extent is not None and extent.cx else None,
                    height_emu=int(extent.cy) if extent is not None and extent.cy else None,
                    alignment=alignment,
                )
            )
            image_index += 1

    return images


def build_docx_bytes(blocks: list[DocxBlock]) -> bytes:
    from docx import Document
    from docx.shared import Emu

    document = Document()
    has_content = False

    for block in blocks:
        if isinstance(block, DocxParagraphBlock):
            paragraph = document.add_paragraph(block.text)
            apply_paragraph_style(document, paragraph, block.style, block.style_id)
            apply_paragraph_alignment(paragraph, block.alignment)
            has_content = has_content or bool(block.text.strip())
            continue

        if isinstance(block, DocxTableBlock):
            row_count = len(block.rows)
            col_count = max((len(row) for row in block.rows), default=0)
            if row_count == 0 or col_count == 0:
                continue

            table = document.add_table(rows=row_count, cols=col_count)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            for row_index, row in enumerate(block.rows):
                for col_index, cell in enumerate(row):
                    docx_cell = table.cell(row_index, col_index)
                    docx_cell.text = cell.text
                    if docx_cell.paragraphs:
                        apply_paragraph_alignment(docx_cell.paragraphs[0], cell.alignment)
                    has_content = has_content or bool(cell.text.strip())
            continue

        if isinstance(block, DocxImageBlock):
            image_bytes = decode_data_url(block.data_url)
            if not image_bytes:
                continue

            paragraph = document.add_paragraph()
            apply_paragraph_alignment(paragraph, block.alignment)
            width = Emu(block.width_emu) if block.width_emu else None
            height = Emu(block.height_emu) if block.height_emu else None
            try:
                paragraph.add_run().add_picture(BytesIO(image_bytes), width=width, height=height)
            except Exception:
                paragraph.add_run(f"[Image omitted: {block.alt_text or block.filename}]")
            has_content = True

    if not has_content:
        document.add_paragraph("")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def summarize_docx_blocks(blocks: list[DocxBlock]) -> str:
    paragraph_count = sum(isinstance(block, DocxParagraphBlock) for block in blocks)
    table_count = sum(isinstance(block, DocxTableBlock) for block in blocks)
    image_count = sum(isinstance(block, DocxImageBlock) for block in blocks)
    return f"内容块 {len(blocks)} 个，段落 {paragraph_count} 个，表格 {table_count} 个，图片 {image_count} 个"


def apply_paragraph_style(document, paragraph, style: str | None, style_id: str | None) -> None:
    for style_key in (style_id, style):
        if not style_key:
            continue
        try:
            paragraph.style = document.styles[style_key]
            return
        except Exception:
            pass
        try:
            paragraph.style = style_key
            return
        except Exception:
            pass


def apply_paragraph_alignment(paragraph, alignment: str | None) -> None:
    if not alignment:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    normalized = alignment.lower().replace("_", "-")
    paragraph.alignment = alignment_map.get(normalized)


def decode_data_url(data_url: str) -> bytes:
    try:
        _, encoded = data_url.split(",", 1)
        return b64decode(encoded)
    except Exception:
        return b""


def get_docx_block_text(block: DocxBlock) -> str:
    if isinstance(block, DocxParagraphBlock):
        return block.text

    if isinstance(block, DocxImageBlock):
        return f"[Image: {block.filename}]"

    return "\n".join("\t".join(cell.text for cell in row) for row in block.rows)
